"""
Plugin para soporte de formatos avanzados (MOBI, AZW3, DOCX)
"""
import os
import logging
from typing import List, Tuple
import subprocess

logger = logging.getLogger(__name__)

class AdvancedFormatsPlugin:
    def __init__(self):
        self.name = "advanced_formats"
        self.description = "Soporte para formatos MOBI, AZW3 y DOCX"

    def register_hooks(self, plugin_manager):
        """Registrar hooks para el sistema de plugins"""
        plugin_manager.register_hook('post_translation', self.on_post_translation)

    def on_post_translation(self, user_id: int, titulo: str, capitulos: List[Tuple[str, str]], formato: str):
        """Hook ejecutado después de una traducción completa"""
        try:
            from main import get_user_setting
            extra_formats = get_user_setting(user_id, 'extra_formats', [])

            if 'mobi' in extra_formats:
                self.generar_mobi(titulo, capitulos)
            if 'azw3' in extra_formats:
                self.generar_azw3(titulo, capitulos)
            if 'docx' in extra_formats:
                self.generar_docx(titulo, capitulos)

        except Exception as e:
            logger.error(f"Error en advanced formats plugin: {e}")

    def generar_mobi(self, titulo: str, capitulos: List[Tuple[str, str]]):
        """Generar archivo MOBI usando calibre"""
        try:
            # Primero generar EPUB temporal
            from main import crear_epub
            epub_file = crear_epub(titulo, capitulos)

            # Convertir a MOBI usando ebook-convert (de calibre)
            mobi_file = epub_file.replace('.epub', '.mobi')

            # Verificar si ebook-convert está disponible
            result = subprocess.run(['which', 'ebook-convert'],
                                  capture_output=True, text=True)

            if result.returncode == 0:
                subprocess.run(['ebook-convert', epub_file, mobi_file],
                             check=True, capture_output=True)
                logger.info(f"Archivo MOBI generado: {mobi_file}")
            else:
                logger.warning("ebook-convert no está instalado. Instala Calibre para soporte MOBI.")

        except Exception as e:
            logger.error(f"Error generando MOBI: {e}")

    def generar_azw3(self, titulo: str, capitulos: List[Tuple[str, str]]):
        """Generar archivo AZW3 usando calibre"""
        try:
            # Primero generar EPUB temporal
            from main import crear_epub
            epub_file = crear_epub(titulo, capitulos)

            # Convertir a AZW3
            azw3_file = epub_file.replace('.epub', '.azw3')

            result = subprocess.run(['which', 'ebook-convert'],
                                  capture_output=True, text=True)

            if result.returncode == 0:
                subprocess.run(['ebook-convert', epub_file, azw3_file],
                             check=True, capture_output=True)
                logger.info(f"Archivo AZW3 generado: {azw3_file}")
            else:
                logger.warning("ebook-convert no está instalado. Instala Calibre para soporte AZW3.")

        except Exception as e:
            logger.error(f"Error generando AZW3: {e}")

    def generar_docx(self, titulo: str, capitulos: List[Tuple[str, str]]):
        """Generar archivo DOCX"""
        try:
            from docx import Document

            doc = Document()
            doc.add_heading(titulo, 0)

            for cap_nombre, cap_texto in capitulos:
                doc.add_heading(cap_nombre, level=1)
                # Dividir texto en párrafos
                paragraphs = cap_texto.split('\n\n')
                for para in paragraphs:
                    if para.strip():
                        doc.add_paragraph(para.strip())

            docx_filename = f"{titulo.replace(' ', '_')}_traducido.docx"
            doc.save(docx_filename)
            logger.info(f"Archivo DOCX generado: {docx_filename}")

        except Exception as e:
            logger.error(f"Error generando DOCX: {e}")

    def get_commands(self):
        """Comandos que este plugin agrega"""
        return {
            'formats': self.show_formats,
            'add_format': self.add_format,
            'remove_format': self.remove_format
        }

    async def show_formats(self, update, context):
        """Mostrar formatos extra configurados"""
        user_id = update.message.from_user.id
        from main import get_user_setting

        formats = get_user_setting(user_id, 'extra_formats', [])
        if formats:
            await update.message.reply_text(f"📄 Formatos extra activados: {', '.join(formats)}")
        else:
            await update.message.reply_text("📄 No tienes formatos extra activados.")

    async def add_format(self, update, context):
        """Agregar formato extra"""
        user_id = update.message.from_user.id
        from main import get_user_setting, set_user_setting

        if not context.args:
            await update.message.reply_text("Uso: /add_format <formato>\nFormatos disponibles: mobi, azw3, docx")
            return

        formato = context.args[0].lower()
        formatos_validos = ['mobi', 'azw3', 'docx']

        if formato not in formatos_validos:
            await update.message.reply_text(f"❌ Formato inválido. Disponibles: {', '.join(formatos_validos)}")
            return

        formats = get_user_setting(user_id, 'extra_formats', [])
        if formato not in formats:
            formats.append(formato)
            set_user_setting(user_id, 'extra_formats', formats)
            await update.message.reply_text(f"✅ Formato {formato.upper()} agregado.")
        else:
            await update.message.reply_text(f"⚠️ El formato {formato.upper()} ya está activado.")

    async def remove_format(self, update, context):
        """Remover formato extra"""
        user_id = update.message.from_user.id
        from main import get_user_setting, set_user_setting

        if not context.args:
            await update.message.reply_text("Uso: /remove_format <formato>")
            return

        formato = context.args[0].lower()
        formats = get_user_setting(user_id, 'extra_formats', [])

        if formato in formats:
            formats.remove(formato)
            set_user_setting(user_id, 'extra_formats', formats)
            await update.message.reply_text(f"✅ Formato {formato.upper()} removido.")
        else:
            await update.message.reply_text(f"⚠️ El formato {formato.upper()} no está activado.")
